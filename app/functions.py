import openai
import json
import textwrap
import numpy as np
from pdfreader import SimplePDFViewer, PDFDocument
import io
import pdfreader
<<<<<<< HEAD
import pandas as pd
import docx
import csv

def gpt3_embedding(content, engine='text-embedding-ada-002'):
    response = openai.Embedding.create(input=content, engine=engine)
    vector = response['data'][0]['embedding']
    return vector


=======
import docx

def gpt3_embedding(content, engine='text-embedding-ada-002'):
    response = openai.Embedding.create(input=content,engine=engine)
    vector = response['data'][0]['embedding']
    return vector

>>>>>>> main
def similarity(v1, v2):
    # Calculate the dot product of two vectors - The dot product of two vectors is a scalar value that indicates how similar the two vectors are. The higher the dot product value, the more similar the two vectors are. The function returns this value as the result. 
    return np.dot(v1, v2)

<<<<<<< HEAD

def search_brain(text, data, count=10):
=======
def search_brain(text, data, count=10): 
>>>>>>> main
    # Generate the embedding for the input text
    vector = gpt3_embedding(text)

    # Create an empty list to store the points for each item in data
    points = list()

    # Loop through each item in data
    for i in data:
        # Calculate the similarity points between the embedding of text and the embedding of the current item
        point = similarity(vector, i['vector'])
        # Add the current item's content and its points to the points list
        points.append({'content': i['content'], 'points': point})

    # Sort the points list in descending order based on the points
    ordered = sorted(points, key=lambda d: d['points'], reverse=True)

    # Return the first count items of the sorted list
    return ordered[0:count]

<<<<<<< HEAD

# GPT-3 Function
def gpt_3(prompt):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt,
        temperature=0,
        max_tokens=1000,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
=======
#GPT-3 Function        
def gpt_3 (prompt):
    response = openai.Completion.create(
    model="text-davinci-003",
    prompt=prompt,
    temperature=0,
    max_tokens=1000,
    top_p=1,
    frequency_penalty=0,
    presence_penalty=0
>>>>>>> main
    )
    text = response['choices'][0]['text'].strip()
    return text

<<<<<<< HEAD

=======
>>>>>>> main
def open_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as infile:
        return infile.read()

<<<<<<< HEAD

=======
>>>>>>> main
def fnBuildBrain(api_key, text_data):
    try:
        openai.api_key = api_key
        chunks = textwrap.wrap(text_data, 2700)
<<<<<<< HEAD

        # Initialize an empty list to store the processed information
        result = list()

        # Loop through each chunk and process its contents
        for chunk in chunks:
            # Get the embedding of the chunk using the gpt3_embedding function
            embedding = gpt3_embedding(chunk.encode(encoding='ASCII', errors='ignore').decode())

            # Store the chunk and its embedding in a dictionary
            info = {'content': chunk, 'vector': embedding}

            # Print the processed information for each chunk
            # print(info, '\n\n\n')

            # Append the processed information to the result list
            result.append(info)

        # Write the result list as a JSON file
        # In this case, indent=2 means that each level in the JSON data will be indented by 2 spaces. This makes the output easier to read, as the structure of the JSON data is clearly defined.
        # print(result)
=======
        
        # Initialize an empty list to store the processed information
        result = list()
        
        # Loop through each chunk and process its contents
        for chunk in chunks:
            # Get the embedding of the chunk using the gpt3_embedding function
            embedding = gpt3_embedding(chunk.encode(encoding='ASCII',errors='ignore').decode())
            
            # Store the chunk and its embedding in a dictionary
            info = {'content': chunk, 'vector': embedding}
            
            # Print the processed information for each chunk
            print(info, '\n\n\n')
            
            # Append the processed information to the result list
            result.append(info)
        
        # Write the result list as a JSON file
        #In this case, indent=2 means that each level in the JSON data will be indented by 2 spaces. This makes the output easier to read, as the structure of the JSON data is clearly defined.
        print(result)
>>>>>>> main
        with open('secondbrain.json', 'w') as outfile:
            json.dump(result, outfile, indent=2)
    except:
        return False

<<<<<<< HEAD

=======
>>>>>>> main
def fnAskQuestion(api_key, question, chunk_size=10000):
    try:
        openai.api_key = api_key
        with open('secondbrain.json', 'r') as infile:
            data = json.load(infile)
            results = search_brain(question, data)
            answers = list()

        # Iterate over the results and generate an answer for each one
        for result in results:
            # Format the prompt for the OpenAI API
            prompt = open_file('qsprompt.txt').replace('<<INFO>>', result['content']).replace('<<QS>>', question)

            # Call the OpenAI API to generate an answer to the user's query
            answer = gpt_3(prompt)

            # Print the answer
            print('\n\n', answer)

            # Add the answer to the list of answers
            answers.append(answer)
<<<<<<< HEAD
            # Join all the answers together into a single string
=======
                # Join all the answers together into a single string
>>>>>>> main
        all_answers = '\n\n'.join(answers)

        # Split the answers into smaller chunks, if necessary
        chunks = textwrap.wrap(all_answers, chunk_size)

        # Initialize a list to store the summaries of the answers
        end = list()

        # Generate a summary for each chunk of answers
        for chunk in chunks:
            # Format the prompt for the OpenAI API
            prompt = open_file('sumanswer.txt').replace('<<SUM>>', chunk)

            # Call the OpenAI API to generate a summary of the answers
            summary = gpt_3(prompt)

            # Add the summary to the list of summaries
            end.append(summary)
        output_result = '\n\n=========\n\n\n\n'.join(end)
        return output_result
    except:
<<<<<<< HEAD
        return ""

=======
        return False
>>>>>>> main

def handle_uploaded_file(f):
    filename = f.name
    extension = filename.split('.')[-1].lower()

    result = ""
    if extension == 'txt':
        result = f.read().decode("utf-8")
        return result
    elif extension == 'pdf':
        with io.BytesIO(f.read()) as f:
            pdf_doc = PDFDocument(f)
            viewer = SimplePDFViewer(f)
            viewer.render()

            # Extract text
            text = ''.join(viewer.canvas.strings)
            return text
    elif extension == "docx":
        # Handle other file types
        document = docx.Document(f)
        full_text = ""
        for para in document.paragraphs:
<<<<<<< HEAD
            full_text = full_text + para.text + "\n"
        return full_text
    elif extension == "xlsx":
        data = pd.read_excel(f.read())
        row_string = ""
        for index, row in data.iterrows():
            # Convert the row to a string
            row_string = ' '.join([str(cell) for cell in row])
            # Print the row string
        print(row_string)
        return row_string
    elif extension == 'csv':
        csv_reader = csv.reader(f.read().decode('utf-8').splitlines())
        row_string = ""
        for row in csv_reader:
            row_string = ' '.join([str(cell) for cell in row])

        print(row_string)
        return row_string

=======
            full_text= full_text + para.text + "\n"
        
        return full_text
>>>>>>> main
